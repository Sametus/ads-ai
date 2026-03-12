import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Hata: {md_path} bulunamadı.")
        return

    doc = Document()
    
    # Başlık stili ayarları (isteğe bağlı olarak özelleştirilebilir)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        
        if not line:
            continue

        # Headers
        header_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            doc.add_heading(text, level=level-1)
            continue

        # Images: ![alt](path)
        image_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if image_match:
            alt_text = image_match.group(1)
            img_path = image_match.group(2)
            
            # Absolute path check for safety
            full_img_path = os.path.abspath(os.path.join(os.path.dirname(md_path), img_path))
            
            if os.path.exists(full_img_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(full_img_path, width=Inches(5.5))
                    
                    # Alt text as caption
                    caption = doc.add_paragraph(alt_text)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f"[Resim Ekleme Hatası: {img_path} - {str(e)}]")
            else:
                doc.add_paragraph(f"[Resim Bulunamadı: {img_path}]")
            continue

        # Lists (Bullets)
        bullet_match = re.match(r'^[\*\-]\s+(.*)', line)
        if bullet_match:
            doc.add_paragraph(bullet_match.group(1), style='List Bullet')
            continue

        # Lists (Numbered)
        numbered_match = re.match(r'^\d+\.\s+(.*)', line)
        if numbered_match:
            doc.add_paragraph(numbered_match.group(1), style='List Number')
            continue

        # Horizontal Rule
        if line == '---':
            # docx'te yatay çizgi için farklı yöntemler var, basitçe boşluk bırakalım
            doc.add_paragraph("_" * 50)
            continue

        # Normal Paragraph
        # Markdown linklerini temizle [text](url) -> text
        processed_line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', line)
        doc.add_paragraph(processed_line)

    doc.save(docx_path)
    print(f"Başarıyla kaydedildi: {docx_path}")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    readme_path = os.path.join(base_dir, "README.md")
    output_path = os.path.join(base_dir, "rapor", "Proje_Raporu_Final.docx")
    
    # Rapor klasörü yoksa oluştur
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    convert_md_to_docx(readme_path, output_path)
